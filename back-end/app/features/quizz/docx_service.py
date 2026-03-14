"""
DOCX Service - Extract text from DOCX files
"""

import os
import logging
import docx

logger = logging.getLogger(__name__)

def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Args:
        docx_path: Path to the DOCX file
    
    Returns:
        Extracted text as a single string
    """
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    all_text = []
    try:
        doc = docx.Document(docx_path)
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text.strip())
                
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() and cell.text.strip() not in all_text:
                        all_text.append(cell.text.strip())
                        
    except Exception as e:
        logger.error(f"Error reading DOCX {docx_path}: {e}")
        raise

    combined = "\n\n".join(all_text)
    logger.info(f"Extracted {len(combined)} chars of text from {os.path.basename(docx_path)}")
    return combined
